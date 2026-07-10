#!/usr/bin/env python3
"""Generate Laporan-Project.docx for aplikasi-penalty (matches Project_Task.pdf criteria)."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Base font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def h1(t):
    p = doc.add_heading(t, level=1)
    return p

def h2(t):
    return doc.add_heading(t, level=2)

def para(t):
    return doc.add_paragraph(t)

def bullet(t):
    return doc.add_paragraph(t, style='List Bullet')

def code(t):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    return p

import base64, subprocess, os, tempfile
def add_mermaid(src, width_in=6.3):
    """Render mermaid source to PNG via mermaid.ink and embed as picture."""
    enc = base64.urlsafe_b64encode(src.encode('utf-8')).decode('ascii')
    url = f'https://mermaid.ink/img/{enc}'
    fd, png = tempfile.mkstemp(suffix='.png')
    os.close(fd)
    try:
        subprocess.run(['curl', '-s', '-H', 'Accept: image/png', '-o', png, url],
                       check=True, timeout=30)
        if os.path.getsize(png) < 500:
            raise RuntimeError('mermaid render too small')
        from docx.shared import Inches
        doc.add_picture(png, width=Inches(width_in))
    finally:
        try: os.remove(png)
        except OSError: pass

# ---- Title ----
title = doc.add_heading('LAPORAN PROJECT', level=0)
sub = doc.add_paragraph('Aplikasi Perhitungan Hasil Lomba Tendangan Penalti')
sub.runs[0].bold = True
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta = doc.add_paragraph('Bahasa C — Arsitektur AES (Agentic Engineering System)')
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# ---- 1. Analisis Permasalahan ----
h1('1. Analisis Permasalahan')

para('Berdasarkan kerangka epistemologis filsafat masalah, suatu kondisi hanya dapat '
     'disebut "masalah" apabila memenuhi tiga prasyarat sekaligus: ')
para('(a) terdapat subjek yang memiliki tujuan; ')
para('(b) terdapat kesenjangan  antara keadaan aktual dan keadaan yang dikehendaki;  ')
para('(c) terdapat hambatan yang tidak dapat diatasi secara langsung. ')

h2('1.1 Tujuan')

para('Membangun program bahasa C untuk mengelola hasil lomba tendangan penalti.')
para('Jumlah peserta 5-7 orang; setiap peserta 7 tendangan; zona bernilai 0 sampai 5.')
para('Mengubah zona menjadi skor; total skor = jumlah seluruh skor tendangan.')
para('Menentukan juara dari total skor tertinggi, dengan aturan seri 5 -> 4 -> 3 -> 2 -> 1.')
para('Menyimpan data peserta, mencatat hasil tendangan, mencari peserta, menampilkan '
     'rekap, dan mengurutkan ranking.')

h2('1.2 Keadaan Saat Ini dan Keadaan Diinginkan')

para('Keadaan saat ini: belum tersedia program yang memenuhi seluruh ketentuan di atas. '
     'Keadaan diinginkan  program yang secara utuh memenuhi batas peserta (5-7), batas '
     'tendangan (7), rentang zona (0-5), aturan konversi dan akumulasi skor, aturan seri '
     'bertingkat, serta kelima kemampuan kelola data. Kesenjangan antara kedua keadaan '
     'inilah yang menjadikan tugas ini sebagai masalah.')

h2('1.3 Hambatan')

bullet('Masalah 1 - Pembatasan input zona 0-5 (Ketentuan 2 & 4). Tanpa pembatasan, '
       'pengguna dapat memasukkan nilai di luar rentang yang merusak perhitungan')
bullet('Masalah 2 - Batas jumlah peserta 5-7 (aturan lomba). Data peserta harus '
       'dialokasikan menampung maksimal 7 tanpa meluap, namun tetap menerima minimal 5.')
bullet('Masalah 3 - Konversi zona ke skor dan akumulasi (Ketentuan 3 & 4). Setiap zona '
       'harus dipetakan ke poin, lalu dijumlahkan menjadi total skor tiap peserta.')
bullet('Masalah 4 - Penentuan juara dan aturan seri bertingkat (Ketentuan 5 & 6). '
       'Pengurutan tidak cukup hanya by total skor; diperlukan pemecah seri 5 -> 4 -> 3 '
       '-> 2 -> 1, dan peringkat sama bila seluruh komponen identik.')
bullet('Masalah 5 - Kelola data antar-fitur (simpan, catat, cari, rekap, ranking). '
       'Seluruh fitur harus beroperasi pada satu kesatuan data peserta yang konsisten.')

# ---- 2. Skenario Program ----
h1('2. Skenario Program')
para('Alur penggunaan aplikasi dalam satu sesi:')
bullet('Menu utama menampilkan 6 fitur + 1 fitur simpan/muat, dengan status tiap '
       'fitur (Aktif / Terkunci / Selesai) sesuai tahap lomba.')
bullet('Tahap 1 — Pendaftaran: pengguna memasukkan nama peserta (5–7). Tombol '
       'peserta ke-8 ditolak; nama kosong mengakhiri pendaftaran.')
bullet('Tahap 2 — Input Tendangan & Skor: untuk tiap peserta, masukkan 7 zona (0–5). '
       'Zona divalidasi; total skor diakumulasi otomatis.')
bullet('Tahap 3 — Ranking & Rekap: setelah semua peserta selesai, pengguna dapat '
       'melihat peringkat (dengan aturan seri) dan rekapitulasi lengkap.')
bullet('Fitur Cari Peserta: mencari peserta berdasarkan nama (cocok persis).')
bullet('Fitur Simpan / Muat Data: menyimpan seluruh lomba ke file data_lomba.bin, '
       'memuatnya kembali saat startup, menghapus file, atau me-reset lomba dari memori.')
para('Verifikasi alur penuh (daftar → tendang → ranking → cari → rekap) telah '
     'ditutupi oleh test otomatis test_full_game_via_surfaces.')

# ---- 3. Konstruksi Program ----
h1('3. Konstruksi Program')
para('Arsitektur yang digunakan adalah AES (Agentic Engineering System) — pola '
     'berlapis ketat (strict layered) dengan dependency inversion dilakukan lewat '
     'struct of function pointers sebagai pengganti interface. Arah dependensi '
     'downward-only: taxonomy -> contract -> capabilities/infrastructure -> agent '
     '-> surfaces -> root (wiring only). Capabilities dan Infrastructure adalah '
     'layer setara (peer) yang sama-sama bergantung ke bawah pada Contract, dan '
     'tidak saling mengimpor.')

para('Hierarki layer (arah panah = arah import yang diizinkan):')
add_mermaid('''graph TD
    subgraph ROOT["root_  ── Wiring Layer (wraps all layers)"]
        direction TB

        S["surface_<br/>(CLI, MCP Server, API)"]
        A["agent_<br/>(Orchestrators)"]

        subgraph PEER["Peer Layers (no direct sibling import)"]
            direction LR
            C["capabilities_<br/>(Checkers, Analyzers)"]
            I["infrastructure_<br/>(Adapters, Scanners)"]
        end

        CT["contract_<br/>(Ports, Protocols, Aggregates)"]
        T["taxonomy_<br/>(VOs, Entities, Errors, Events, Constants)"]

        S -->|"imports"| CT
        S -->|"imports"| T
        A -->|"imports"| CT
        A -->|"imports"| T
        C -->|"imports"| CT
        C -->|"imports"| T
        I -->|"imports"| CT
        I -->|"imports"| T
        CT -->|"imports"| T
    end

    ROOT_CONT["root_container<br/>(DI Wiring — instantiates & injects all)"]
    ROOT_ENTRY["root_entry<br/>(Binary Bootstrap)"]

    ROOT_CONT -->|"wires"| ROOT
    ROOT_ENTRY -->|"starts"| ROOT_CONT''')
para('Alur data utama aplikasi:')
add_mermaid('''graph LR
A[Menu Utama] --> B[Pendaftaran 5-7 Peserta]
B --> C[Input Tendangan: 7 zona 0-5]
C --> D[Konversi Zona ke Skor]
D --> E{Total Skor Sama?}
E -->|Ya| F[Pemecah Seri 5-4-3-2-1]
E -->|Tidak| G[Ranking Final]
F --> G
G --> H[Rekap dan Cari Peserta]
H --> I[Simpan / Muat File]''')

# ---- 4. Struktur (struct) ----
h1('4. Struktur (struct)')
para('Struct utama (diekstrak dari source):')
h2('CompetitionState (wadah status lomba)')
code('typedef struct {\n'
     '    ParticipantEntity     participants[MAX_PARTICIPANTS];\n'
     '    int                   participant_count;\n'
     '    CompetitionStateKind  state;   /* INIT | REGISTERED | COMPLETED */\n'
     '} CompetitionState;')
h2('ParticipantEntity (satu peserta)')
code('typedef struct {\n'
     '    ParticipantIdVO   id;          /* nomor urut */\n'
     '    ParticipantNameVO name;        /* nama (dibungkus VO) */\n'
     '    KickVO            kicks[TOTAL_KICKS]; /* 7 hasil tendangan */\n'
     '    TotalScoreVO      total_score; /* akumulasi poin */\n'
     '    ZoneFreqVO        zone_freq;   /* hitungan tiap zona (pemecah seri) */\n'
     '    KickCountVO       kick_count;  /* 0..7 tendangan dilakukan */\n'
     '} ParticipantEntity;')
h2('CompetitionStateKind (enum tahap)')
code('typedef enum {\n'
     '    STATE_INIT = 0,       /* belum ada peserta */\n'
     '    STATE_REGISTERED = 1, /* boleh tendang & cari */\n'
     '    STATE_COMPLETED = 2   /* boleh ranking & recap */\n'
     '} CompetitionStateKind;')

# ---- 5. Konstanta ----
h1('5. Konstanta')
para('Konstanta terpusat di taxonomy_game_constant.h:')
code('MIN_PARTICIPANTS 5     /* peserta minimal */\n'
     'MAX_PARTICIPANTS 7     /* peserta maksimal (batas array) */\n'
     'TOTAL_KICKS      7     /* tendangan per peserta */\n'
     'MIN_ZONE         0     /* zona terendah (miss) */\n'
     'MAX_ZONE         5     /* zona tertinggi (top) */\n'
     'MAX_NAME_LENGTH  30    /* panjang nama maksimal */\n'
     'DEFAULT_STORAGE_FILENAME "data_lomba.bin"')

# ---- 6. Variabel ----
h1('6. Variabel')
para('Aplikasi sengaja TIDAK menggunakan variabel global untuk data lomba. '
     'Variabel utama:')
bullet('CompetitionState state — dideklarasikan di main(), di-pass ke seluruh fitur '
       'via pointer (satu sumber data).')
bullet('Aggregate lokal di main(): reg, sc, rk, sr, rc, st, sn — masing-masing '
       'berisi function pointer ke implementasi domain.')
bullet('DisplayPort dp — antarmuka render, dirakit oleh root_display_build().')
bullet('Variabel lokal di tiap fungsi: loop index (i, z), buffer (buf[128]), '
       'state sementara (RankingEntryVO entries[]).')

# ---- 7. Fungsi ----
h1('7. Daftar Fungsi')
para('Fungsi domain & infrastruktur utama:')
bullet('root_registration_build / agent_registration_append (pendaftaran peserta)')
bullet('capabilities_scoring_validate_zone / capabilities_scoring_record_kick (input tendangan & skor)')
bullet('capabilities_ranking_compute (urutkan + aturan seri zona 5→4→3→2→1)')
bullet('capabilities_search_resolver (cari peserta)')
bullet('agent_recap_orchestrator / capabilities_recap_formatter (rekapitulasi)')
bullet('agent_storage_save / agent_storage_load / agent_storage_delete (simpan/muat/hapus file)')
bullet('sanitizer_validate_int / sanitizer_validate_string (validasi input)')
bullet('cli_surfaces_menu_run + cli_surfaces_*_execute (navigasi & layar tiap fitur)')
bullet('root_display_build (rakit DisplayPort / antarmuka ncurses)')
bullet('cli_surfaces_storage_execute (fitur Simpan/Muat/Reset baru)')

# ---- 8. Kode Sumber ----
h1('8. Kode Sumber (Script Program)')
para('Kode sumber lengkap terdapat di folder src/ (41 file .c/.h) dengan struktur:')
bullet('src/shared/        — konstanta, struct (VO), enum, kontrak antarmuka')
bullet('src/registration/  — pendaftaran peserta')
bullet('src/scoring/       — validasi & pencatatan zona → skor')
bullet('src/ranking/       — peringkat + aturan seri')
bullet('src/search/        — pencarian peserta')
bullet('src/recap/         — rekapitulasi')
bullet('src/storage/       — simpan/muat/hapus file')
bullet('src/sanitizer/     — validasi input')
bullet('src/cli/           — layar menu & tiap fitur (command + page)')
bullet('src/tui/           — adaptor ncurses (DisplayPort)')
bullet('root_cli_main_entry.c — titik masuk program')
para('Cuplikan fungsi inti — aturan seri (ranking):')
code('static int compare_entries(const void *a, const void *b) {\n'
     '    const RankingEntryVO *x = a, *y = b;\n'
     '    if (x->total_score != y->total_score)\n'
     '        return y->total_score - x->total_score;\n'
     '    for (int z = MAX_ZONE; z >= 1; z--)        /* 5,4,3,2,1 */\n'
     '        if (x->zone_freq[z] != y->zone_freq[z])\n'
     '            return y->zone_freq[z] - x->zone_freq[z];\n'
     '    return 0;                                  /* seri sempurna */\n'
     '}')
para('Kompilasi & pengujian: make (build) dan make test (semua test lolos).')

# ---- Header identitas: HANYA halaman pertama ----
# Tidak dibentuk sendiri; diambil dari identitas penyusun (template tugas).
from docx.oxml.ns import qn

sec = doc.sections[0]
sectPr = sec._sectPr

# Aktifkan "Different First Page"
titlePg = sectPr.find(qn('w:titlePg'))
if titlePg is None:
    titlePg = sectPr.makeelement(qn('w:titlePg'), {})
    sectPr.append(titlePg)
titlePg.set(qn('w:val'), 'true')

# Isi first-page header dengan identitas
first_hdr = sec.first_page_header
first_hdr.is_linked_to_previous = False
# Hapus isi lama (jika ada)
for p in list(first_hdr.paragraphs):
    p._p.getparent().remove(p._p)

# Header 2 kolom: kiri (Nama, Nim), kanan (Mata Kuliah/Kelas, Sesi)
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT

from docx.shared import Inches
tbl = first_hdr.add_table(rows=2, cols=2, width=Inches(6.5))
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
tbl.autofit = True
cells = tbl.rows[0].cells + tbl.rows[1].cells
left = [('Nama', 'Raka Arwaky'), ('Nim', '22342030')]
right = [('Mata Kuliah', 'Pengantar Coding'), ('Sesi', '863')]
for i in range(2):
    lc = cells[i]; rc = cells[i + 2]
    lc.text = ''; rc.text = ''
    lc.paragraphs[0].add_run(f'{left[i][0]} : {left[i][1]}')
    rc.paragraphs[0].add_run(f'{right[i][0]} : {right[i][1]}')
# Hapus border tabel biar tampak seperti teks biasa
tblPr = tbl._tbl.tblPr
borders = tblPr.makeelement(qn('w:tblBorders'), {})
for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
    e = borders.makeelement(qn('w:' + edge), {qn('w:val'): 'none', qn('w:sz'): '0'})
    borders.append(e)
tblPr.append(borders)

# Pastikan tidak ada default header (biar halaman 2+ kosong)
sec.header.is_linked_to_previous = True

doc.save('Laporan-Project.docx')
print('SAVED Laporan-Project.docx')
